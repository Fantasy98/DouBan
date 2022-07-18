from ast import Break, keyword
from base64 import encode
from email import header
from os import link


import requests 
from bs4 import BeautifulSoup 
from docx  import Document 
from docx import shared
from docx.shared import Inches,Pt
import os


def getLink(keyword,pagenumber):
    url =(u"https://book.douban.com/latest?subcat={}&p={}".format(keyword,pagenumber))
    head = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/100.0.4896.127 Safari/537.36 Edg/100.0.1185.44'}

    try: 
            req = requests.get(url,headers=head)
    except requests.exceptions.RequestException:
            return None
    return BeautifulSoup(req.text,'lxml')


def getNames(bs):
    Namelist = []
   
    info = bs.find_all('a',{'class':'fleft'})
    for i in info: 
        Namelist.append(list(i)[0])
    return Namelist


def getLinks(bs):
    Linklist = []
   
    info = bs.find_all('a',{'class':'fleft'})
    for i in info: 
        Linklist.append(i['href'])
        
    return Linklist



######以下是获取一本书的所有信息

#获取 bs
def getABook(url): 
    head = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/100.0.4896.127 Safari/537.36 Edg/100.0.1185.44'}

    try: 
            req = requests.get(url,headers=head)
    except requests.exceptions.RequestException:
            return None
    return BeautifulSoup(req.text,'lxml')

#获取定价，作者，出版社等信息
def getABasic(bs):
    basic = ''
    info = bs.find('div',{'id':"info", 'class':""})
    for i in info.get_text().split():
         basic+= i + '\n'
        
    return basic

#获取评分和打分占比
def getMark(bs):
    mark = bs.find('div', {'id':"interest_sectl", 'class':""}) 
    Score=mark.find('strong', {'class':"ll rating_num" ,'property':"v:average"}).get_text()# 豆瓣评分
    return Score

def getPercent(bs):
    Percentage =[]
    mark = bs.find('div', {'id':"interest_sectl", 'class':""})
    percentage = mark.find_all('span',{'class':"rating_per"})
    for per in percentage: 
        Percentage.append(per.get_text())
    return Percentage

#作者简介
def getAuthorIntro(bs):
    Author = bs.find_all('div',{'class':"intro"})
    if len(Author)>=2:
        if Author[1] is not None:
            
                Auth = Author[1].get_text()
                return Auth
        else:
                return Author[0].get_text()
    else:
        return None
    

#书内容简介
def getBookIntro(bs):
    Intro = bs.find_all('div',{'class':"intro"})
    BookIntro = Intro[-1].get_text()
    return BookIntro

#封面获取
def getCover(bs):
    head = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/100.0.4896.127 Safari/537.36 Edg/100.0.1185.44'}
    cover = bs.find('a', {'class':"nbg"})
    jpg = requests.get ( cover['href'],headers=head)
    return jpg


class Onebook:
    def __init__(self,url):
        self.url =url
        self.bs= getABook(self.url)
        self.info = getABasic(self.bs) #基本信息
        self.mark= getMark(self.bs) #评分和评分占比
        self.distrib = getPercent(self.bs)
        self.author = getAuthorIntro(self.bs)#作者简介
        self.Intro = getBookIntro(self.bs)#书内容简介
        self.cover = getCover(self.bs)
    


keywords = {1:'全部',2:'文学',3:'小说',4:'历史文化',5:'社会纪实',6:'科学新知',7:'艺术设计',8:'商业经管',9:'绘本漫画'}
print('今天想看什么？\n1\tAll \n2\t文学\n3\t小说\n4\t历史文化\n5\t社会纪实\n6\t科学新知\n7\t艺术设计\n8\t商业经管\n9\t绘本漫画')

Key = int(input())

key = keywords[Key]

doc1 = Document()
Head1 = doc1.add_heading("{}新书推荐".format(key),1)
for page in range(0,10):
    
    Name=getNames(getLink(keyword=key,pagenumber=page+1))
    Link = getLinks(getLink(keyword=key,pagenumber=page+1))
    if len(Name)==0:
        print('All Over !')
    Break

    print("At Page{},{} items totally".format(page+1,len(Name)))
    
    for i in range(len(Name)): 

        print('Writing\t{}'.format(Name[i]))
        print('Link\t{}'.format(Link[i]))
        h1 =doc1.add_heading(u"{}".format(Name[i]),2)

        Abook = Onebook(url=Link[i])

        with open('{}.jpg'.format(i),'wb') as f : 
            f.write(Abook.cover.content)
        doc1.add_picture("{}.jpg".format(i),width=shared.Cm(6))

        h2 =doc1.add_heading('基本信息',3)
        p1 =doc1.add_paragraph(u"{}".format(Abook.info))
        

        h3 =doc1.add_heading('豆瓣评分',3)
        p2=doc1.add_paragraph(u"{}".format(Abook.mark))
        

        h4 =doc1.add_heading('评分占比',3)
        if Abook.distrib is not None and len(Abook.distrib) != 0:
            p3 =doc1.add_paragraph('五星：\t{}\n 四星：\t{}\n 三星：\t{}\n 二星：\t{}\n 一星：\t{}\n'.format(Abook.distrib[0],Abook.distrib[1],Abook.distrib[2],Abook.distrib[3],Abook.distrib[4]))
        else:
            p3=doc1.add_paragraph('暂无')

        h5=doc1.add_heading('作者简介',3)
        p4=doc1.add_paragraph(u"{}".format(Abook.Intro))
        p4.paragraph_format.first_line_indent = Inches(1)

        h6 =doc1.add_heading('作品简介',3)
        p5=doc1.add_paragraph(u"{}".format(Abook.author))
        p5.paragraph_format.first_line_indent = Inches(1)

        doc1.add_page_break()  
        rootdir = r"/Users/wangyuning/Desktop/DouBan/"
        filelist = os.listdir(rootdir)
        for file in filelist:
            if '.jpg' in file : 
                del_file =rootdir + '//'+file
                os.remove(del_file)


import datetime

doc1.save(u'{}新书推荐{}.docx'.format(str(key),datetime.date.today()))

        


        

