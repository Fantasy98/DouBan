
from operator import index

from typing import Dict
from numpy import sort
import requests
from bs4 import BeautifulSoup 
import json 
from docx import Document
from docx import shared
from docx.shared import Inches,Pt




#For a tv 

def Getit(url):
        header1 = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/100.0.4896.127 Safari/537.36 Edg/100.0.1185.44'}
        try:
            req = requests.get(url,headers=header1)
        except requests.exceptions.RequestException:
                return None
        return  BeautifulSoup(req.text,'lxml')

def GetTitle(bs):
    title  = bs.find('meta', {'property':"og:title"})['content']
    return title

def GetInfo(bs): 
        info = bs.find('div', {'id':"info"})
        information = info.get_text()
        return information

def GetScore(bs):
        score = bs.find('strong',{ 'class':"ll rating_num" ,'property':"v:average"})
        Socre = score.get_text()   
        return float(Socre)

def Getdistri(bs):
    #Rating distribution
    Distribution = []
    distri = bs.find_all('span', {'class':"rating_per"})
    for i in distri:
        Distribution.append(i.get_text())
    return Distribution

def GetComment(bs):
    comment = bs.find('span',{ 'property':"v:votes"})
    if comment is not None:
        Comment = comment.get_text()
        return float(Comment)

def GetActors(bs):
    ACTORS =[]
    actors =bs.find_all('meta',{'property':"video:actor"})
    for i in actors[0:6]: 
       ACTORS.append( i['content'])

def GetCoverurl(bs):
    header1 = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/100.0.4896.127 Safari/537.36 Edg/100.0.1185.44'}
    cover = bs.find('span',{'class': "rec"}).find('a',{'class':'bn-sharing'})
    url= cover['data-pic']
    Cover = requests.get(url,headers=header1)
    return Cover



class ATV:
    def __init__(self,url):
        self.url = url
        self.bs = Getit(self.url)
        self.Name = GetTitle(self.bs)
        self.info = GetInfo(self.bs)
        self.socre = GetScore(self.bs)
        self.distri = Getdistri(self.bs)
        self.Comments  = GetComment(self.bs)
        self.Actors = GetActors(self.bs)
        self.Cover = GetCoverurl(self.bs)


#######对抓包素材预处理
def Package(url): 
    header1 = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/100.0.4896.127 Safari/537.36 Edg/100.0.1185.44'}
    try:
            req = requests.get(url,headers=header1)
    except requests.exceptions.RequestException:
                return None
    return  req.content.decode()

def Urls(json_text):
    Dictionary = json.loads(json_text)
    b = Dictionary['subjects']
    Links = [i['url'] for i in b]
    return Links

def Titles(json_text):
    Dictionary = json.loads(json_text)
    b = Dictionary['subjects']
    titles = [i['title'] for i in b]
    return titles



base= 20 
Total = []
Names =[]
for i in range(0,6):
    step = base*i
    url ="https://movie.douban.com/j/search_subjects?type=tv&tag=美剧&sort=recommend&page_limit=20&page_start={}".format(step)
    json_text = Package(url)
    Links = Urls(json_text=json_text)
    # titles = Titles(json_text=json_text) 
    # j=0
    for link in Links : 
        # j+=1
        print('checking\t{}'.format(link))
        bs = Getit(link)
        comment_number =GetComment(bs)
       
        print('The comment number is\t{}'.format(comment_number))
        if (comment_number is None)  or (comment_number <10000) : 
            # titles.remove(titles[i-1])

            Links.remove(link)
            
            print('{} has been moved!'.format(link))
    Total+=Links
    # Names+=titles



#录入所有链接的信息为dictionary


#写文档

i = 0
Tv_Dict = {}# Empty Dictionary for each Tv
TVs = [] # List for all the TVs for sake of sorted 
for link in Total:
    i+=1
   
    Tv = ATV(link)
    Tv_Dict.update({'Name':Tv.Name})
    print( u'Updating\t{}'.format(Tv_Dict['Name']))
    Tv_Dict.update({'Link':link})
    Tv_Dict.update({'Comments':Tv.Comments})
    Tv_Dict.update({'Score':Tv.socre})
    Tv_Dict.update({'Distribution':Tv.distri})
    Tv_Dict.update({'Info':Tv.info})
    TVs.append(Tv_Dict)
    Tv_Dict = {}

TVs_sorted = sorted(TVs ,key= lambda n : n['Score'],reverse=1)
print("Sortting Complete!")


doc = Document()
doc.add_heading('豆瓣热门美剧【按评论筛选】',1)
i= -1
for tv in TVs_sorted:
    i +=1
    print(u'Wrtting\t{}'.format(tv['Name']))

    doc.add_heading('{}'.format(tv['Name']))
    
    bs = Getit(tv['Link'])
    jpeg = GetCoverurl(bs)

    with open('{}.jpeg'.format(i),'wb') as f : 
            f.write(jpeg.content)

    doc.add_picture("{}.jpeg".format(i),width=shared.Cm(6))
    doc.add_paragraph('{}'.format(tv['Link']))
    
    doc.add_heading('评论人数',3)
    doc.add_paragraph('{} 条'.format(tv['Comments']))
    doc.add_heading('豆瓣评分',3)
    doc.add_paragraph('{}'.format(tv['Score'],'0.1f'))
    doc.add_heading('评分分布',3)
    if tv['Distribution'] is not None:
            p3 =doc.add_paragraph('五星：\t{}\n 四星：\t{}\n 三星：\t{}\n 二星：\t{}\n 一星：\t{}\n'.format(tv['Distribution'][0],tv['Distribution'][1],tv['Distribution'][2],tv['Distribution'][3],tv['Distribution'][4]))
    else:
            p3=doc.add_paragraph('暂无')
    doc.add_heading('剧集信息',3)
    doc.add_paragraph('{}'.format(tv['Info']))

    doc.add_page_break()
    import os 
    rootdir = r"C:\Users\王裕宁\Desktop\DOUBAN TV"
    filelist = os.listdir(rootdir)
    for file in filelist:
        if '.jpeg' in file : 
            del_file =rootdir + '//'+file
            os.remove(del_file)


import datetime

doc.save(u'美剧{}.docx'.format(datetime.date.today()))



